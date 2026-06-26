import streamlit as st
import os
import json
import bcrypt
import smtplib
import ssl
import random
import string
import uuid
import requests
import io
import urllib3
from openai import OpenAI
from email.message import EmailMessage
from datetime import datetime, timedelta
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import streamlit as st

# ==================== КОНФИГУРАЦИЯ ====================
USERS_PATH = "data/users.json"
PENDING_PATH = "data/pending_users.json"
os.makedirs("data", exist_ok=True)

# Загрузка переменных: сначала из st.secrets (Streamlit Cloud), потом из .env (локально)
DEEPSEEK_API_KEY = st.secrets.get("DEEPSEEK_API_KEY") or os.getenv("DEEPSEEK_API_KEY")
GIGACHAT_CLIENT_ID = st.secrets.get("GIGACHAT_CLIENT_ID") or os.getenv("GIGACHAT_CLIENT_ID")
GIGACHAT_CLIENT_SECRET = st.secrets.get("GIGACHAT_CLIENT_SECRET") or os.getenv("GIGACHAT_CLIENT_SECRET")
SMTP_HOST = st.secrets.get("SMTP_HOST") or os.getenv("SMTP_HOST", "smtp.yandex.ru")
SMTP_PORT = int(st.secrets.get("SMTP_PORT") or os.getenv("SMTP_PORT", "465"))
SMTP_USER = st.secrets.get("SMTP_USER") or os.getenv("SMTP_USER")
SMTP_PASS = st.secrets.get("SMTP_PASS") or os.getenv("SMTP_PASS")

# ==================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ====================
def load_json(path):
    if not os.path.exists(path): return {}
    with open(path, "r", encoding="utf-8") as f: return json.load(f)

def save_json(path, data):
    with open(path, "w", encoding="utf-8") as f: json.dump(data, f, ensure_ascii=False, indent=2)

def hash_pwd(pwd): return bcrypt.hashpw(pwd.encode(), bcrypt.gensalt()).decode()
def check_pwd(pwd, hashed): return bcrypt.checkpw(pwd.encode(), hashed.encode())

def send_code_email(email, code):
    msg = EmailMessage()
    msg["Subject"] = "🔐 Код подтверждения: Ассистент ДОУ"
    msg["From"] = SMTP_USER
    msg["To"] = email
    msg.set_content(f"Ваш код: {code}\n⏳ Действует 15 минут.\n🔒 Не передавайте третьим лицам.")
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, context=ctx) as srv:
            srv.login(SMTP_USER, SMTP_PASS)
            srv.send_message(msg)
        return True
    except Exception as e:
        st.error(f"⚠️ Ошибка отправки email: {str(e)}")
        return False

# ==================== ЭКСПОРТ И ИЗОБРАЖЕНИЯ ====================
def export_to_pdf(text, title="Документ ДОУ"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, title, align="C")
    pdf.ln(10)
    pdf.set_font("DejaVu", "", 12)
    for line in text.split("\n"):
        if line.strip():
            pdf.multi_cell(0, 6, line)
            pdf.ln(2)
    return pdf.output(dest="S").encode("latin-1") if pdf.output else pdf.output()

def export_to_docx(text, title="Документ ДОУ"):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_image_shevdevrum(prompt):
    try:
        safe_prompt = requests.utils.quote(prompt.replace(" ", "+"))
        img_url = f"https://image.pollinations.ai/prompt/{safe_prompt}?width=1024&height=768&nologo=true"
        return img_url
    except Exception as e:
        st.error(f"⚠️ Ошибка генерации изображения: {str(e)}")
        return None

# ==================== ИИ-ДВИЖОК ====================
DEEPSEEK_CLIENT = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

@st.cache_resource(ttl=1700)
def get_gigachat_token():
    headers = {
        "Authorization": f"Basic {GIGACHAT_AUTH_KEY}",
        "Content-Type": "application/x-www-form-urlencoded",
        "Accept": "application/json",
        "RqUID": str(uuid.uuid4())
    }
    data = {"scope": "GIGACHAT_API_PERS"}
    resp = requests.post(
        "https://ngw.devices.sberbank.ru:9443/api/v2/oauth",
        headers=headers,
        data=data,
        verify=False
    )
    resp.raise_for_status()
    return resp.json()["access_token"]

def ask_ai(prompt, model, age_group, use_kustuk=False):
    system = """Ты методический ИИ-ассистент для воспитателей ДОУ. Работай строго по ФОП ДО, ФГОС ДО, СанПиН 2.4.3648-20 и 273-ФЗ."""
    
    if use_kustuk:
        system += """
        
🌈 РЕГИОНАЛЬНАЯ ПРОГРАММА «КУСТУК» (Республика Саха (Якутия)):
- Интегрируй этнокультурный компонент: якутский фольклор, традиции, природу Севера
- Включай национальные подвижные и дидактические игры
- Учитывай художественно-эстетическое развитие через олонхо, тойук, якутский орнамент
- Физическое воспитание: традиционные упражнения
- Речевое развитие: элементы якутского языка (посильно возрасту)
- Социально-коммуникативное: уважение к старшим, бережное отношение к природе
- При генерации указывай, какие элементы относятся к федеральному компоненту, а какие — к региональному «Кустук»."""
    
    system += f"""
Возрастная группа: {age_group}. 
⚠️ Не запрашивай и не храни ФИО детей. Используй ID или обобщённые формулировки.
Если норматив не найден — помечай: «Требует методической проверки».
Формат: структурированный текст, готовый к копированию."""

    try:
        if model == "GigaChat":
            token = get_gigachat_token()
            headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json", "RqUID": str(uuid.uuid4())}
            payload = {"model": "GigaChat", "messages": [{"role":"system","content":system},{"role":"user","content":prompt}], "temperature":0.3, "max_tokens":2500}
            resp = requests.post(
                "https://gigachat.devices.sberbank.ru/api/v1/chat/completions", 
                headers=headers, 
                json=payload,
                verify=False
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
        else:
            res = DEEPSEEK_CLIENT.chat.completions.create(model="deepseek-chat", messages=[{"role":"system","content":system},{"role":"user","content":prompt}], temperature=0.3)
            return res.choices[0].message.content
    except Exception as e:
        return f"️ Ошибка ИИ ({model}): {str(e)}"

# ==================== АВТОРИЗАЦИЯ ====================
if "auth_step" not in st.session_state: st.session_state.auth_step = "login"
st.set_page_config(page_title="Ассистент ДОУ «Кустук»", layout="wide")

if st.session_state.auth_step == "login":
    st.title("🔐 Вход в систему")
    email = st.text_input("Email")
    password = st.text_input("Пароль", type="password")
    col1, col2 = st.columns(2)
    if col1.button("Войти", use_container_width=True):
        users = load_json(USERS_PATH)
        u = users.get(email)
        if u and check_pwd(password, u["password"]):
            st.session_state.user = u
            st.session_state.auth_step = "dashboard"
            st.rerun()
        else: st.error("Неверный email или пароль")
    if col2.button("📝 Регистрация", use_container_width=True):
        st.session_state.auth_step = "register"
        st.rerun()
    st.stop()

elif st.session_state.auth_step == "register":
    st.title("📝 Регистрация воспитателя")
    em = st.text_input("Корпоративный email")
    name = st.text_input("ФИО")
    p1 = st.text_input("Пароль", type="password")
    p2 = st.text_input("Повторите пароль", type="password")
    if st.button("Получить код", use_container_width=True):
        if not all([em,name,p1,p2]): st.error("Заполните все поля")
        elif p1!=p2: st.error("Пароли не совпадают")
        elif em in load_json(USERS_PATH) or em in load_json(PENDING_PATH): st.error("Email уже используется")
        else:
            code = "".join(random.choices(string.digits, k=6))
            save_json(PENDING_PATH, load_json(PENDING_PATH) | {em: {"name":name, "password":hash_pwd(p1), "role":"teacher", "code":code, "expires":(datetime.now()+timedelta(minutes=15)).isoformat()}})
            if send_code_email(em, code):
                st.session_state.pending_email = em
                st.session_state.auth_step = "verify"
                st.success("✅ Код отправлен. Проверьте папку «Спам».")
                st.rerun()
    if st.button("← Назад"): st.session_state.auth_step = "login"; st.rerun()
    st.stop()

elif st.session_state.auth_step == "verify":
    st.title("🔑 Подтверждение email")
    st.info(f"Код отправлен на: {st.session_state.pending_email}")
    code_input = st.text_input("6-значный код")
    if st.button("Подтвердить", use_container_width=True):
        pend = load_json(PENDING_PATH)
        usr = pend.get(st.session_state.pending_email)
        if not usr: st.error("Сессия истекла"); st.session_state.auth_step="register"
        elif datetime.now().isoformat() > usr["expires"]: st.error("⏳ Код просрочен")
        elif code_input != usr["code"]: st.error("❌ Неверный код")
        else:
            users = load_json(USERS_PATH)
            users[st.session_state.pending_email] = {"name":usr["name"], "password":usr["password"], "role":usr["role"], "created":datetime.now().isoformat()}
            save_json(USERS_PATH, users)
            del pend[st.session_state.pending_email]
            save_json(PENDING_PATH, pend)
            st.session_state.user = users[st.session_state.pending_email]
            st.session_state.auth_step = "dashboard"
            st.success("✅ Аккаунт активирован!")
            st.rerun()
    st.stop()

# ==================== ГЛАВНЫЙ ИНТЕРФЕЙС ====================
elif st.session_state.auth_step == "dashboard" and "user" in st.session_state:
    user = st.session_state.user
    st.sidebar.success(f"👤 {user['name']} | {user['role'].capitalize()}")
    if st.sidebar.button("🚪 Выйти"):
        del st.session_state.user; st.session_state.auth_step="login"; st.rerun()

    ai_model = st.sidebar.selectbox(" Модель ИИ", ["GigaChat", "DeepSeek"])
    age_group = st.sidebar.selectbox("🎯 Возрастная группа", ["2-3 года", "3-4 года", "4-5 лет", "5-6 лет", "6-7 лет"])
    use_kustuk = st.sidebar.checkbox("🌈 Программа «Кустук»", value=False)
    if use_kustuk:
        st.sidebar.info("Активирован региональный компонент РС(Я)")

    menu = ["📅 Календарный план", "📊 Диагностика", "🤖 ИИ-чат", "🌈 Программа «Кустук»", "📖 НПА"]
    if user["role"] == "admin": menu.append("👥 Пользователи")
    page = st.sidebar.radio("Модули", menu)

    st.title(page)

    if "last_result" not in st.session_state: st.session_state.last_result = ""
    if "last_img" not in st.session_state: st.session_state.last_img = None

    def show_exports(text, title_prefix="Документ"):
        col1, col2 = st.columns(2)
        if col1.button(" Скачать PDF"):
            pdf_bytes = export_to_pdf(text, f"{title_prefix} ({age_group})")
            st.download_button("💾 Скачать PDF", data=pdf_bytes, file_name=f"{title_prefix.lower().replace(' ','_')}.pdf", mime="application/pdf")
        if col2.button("📝 Скачать Word"):
            doc_bytes = export_to_docx(text, f"{title_prefix} ({age_group})")
            st.download_button("💾 Скачать Word", data=doc_bytes, file_name=f"{title_prefix.lower().replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if page == "📅 Календарный план":
        theme = st.text_input("Тема недели/месяца")
        focus = st.multiselect("Фокус развития", ["Речь", "Познание", "Социализация", "Движение", "Творчество"])
        if st.button("🧠 Сгенерировать план"):
            kustuk_note = "\n\nДОБАВЬ региональный компонент «Кустук»: укажи интеграцию якутского фольклора, национальных игр, этнокультурных ценностей. Раздели пункты на федеральные и региональные." if use_kustuk else ""
            with st.spinner("Генерация..."):
                res = ask_ai(f"Составь календарный план для группы {age_group}. Тема: '{theme}'. Фокус: {', '.join(focus)}. Включи режимные моменты, НОД, прогулку, игры, работу с родителями. Соответствие ФОП/ФГОС.{kustuk_note}", ai_model, age_group, use_kustuk)
            st.session_state.last_result = res
            st.markdown(res)
            show_exports(res, "Календарный план")
            st.divider()
            if st.button("🖼️ Сгенерировать иллюстрацию к теме"):
                with st.spinner("Создаю изображение..."):
                    st.session_state.last_img = generate_image_shevdevrum(f"детский сад, образовательная иллюстрация, {theme}, {age_group}, стиль: акварель, добрый, методический материал")
            if st.session_state.last_img:
                st.image(st.session_state.last_img, caption=f"Иллюстрация: {theme}", use_container_width=True)

    elif page == "📊 Диагностика":
        period = st.selectbox("Период", ["Начало года", "Середина года", "Конец года"])
        domain = st.selectbox("Область", ["Речевое", "Познавательное", "Социально-коммуникативное", "Физическое", "Художественно-эстетическое"])
        data_in = st.text_area("Наблюдаемые показатели (без ФИО, используйте ID или обобщённо)")
        if st.button("🔍 Анализ"):
            kustuk_note = "\n\nОцени динамику с учётом регионального компонента «Кустук»." if use_kustuk else ""
            with st.spinner("Анализ..."):
                res = ask_ai(f"Анализ диагностики за {period}, группа {age_group}, область: {domain}. Данные: {data_in}. Выводы, динамика, коррекционные действия. ФГОС/ФОП.{kustuk_note}", ai_model, age_group, use_kustuk)
            st.session_state.last_result = res
            st.markdown(res)
            show_exports(res, "Диагностика")

    elif page == " ИИ-чат":
        if "chat" not in st.session_state: st.session_state.chat = []
        inp = st.chat_input("Задайте вопрос...")
        if inp:
            st.session_state.chat.append({"role":"user","content":inp})
            with st.spinner("Думаю..."):
                full = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.chat[-4:]])
                reply = ask_ai(full, ai_model, age_group, use_kustuk)
                st.session_state.chat.append({"role":"assistant","content":reply})
            for m in st.session_state.chat:
                with st.chat_message(m["role"]): st.write(m["content"])
            st.session_state.last_result = reply
            st.divider()
            show_exports(reply, "Диалог с ИИ")

    elif page == "🌈 Программа «Кустук»":
        st.markdown("### 📜 Региональная программа «Кустук» (РС Якутия)")
        kustuk_section = st.selectbox("📂 Раздел генерации", ["📅 Тематическое планирование", "🎮 Национальные игры", "📚 Фольклор и язык", "🎨 Творчество", "🤝 Работа с родителями"])
        
        prompt_map = {
            "📅 Тематическое планирование": lambda t: f"Тематический план по «Кустук» на тему: '{t}'. Раздели на федеральный и региональный блоки.",
            "🎮 Национальные игры": lambda _: f"Подбери якутские национальные игры для {age_group} по «Кустук». Укажи правила, оборудование, развиваемые качества.",
            "📚 Фольклор и язык": lambda _: f"Подбери якутский фольклор для {age_group} с переводом и методическими рекомендациями.",
            " Творчество": lambda _: f"Конспект занятия по художественному творчеству в этностиле для {age_group} по «Кустук».",
            "🤝 Работа с родителями": lambda _: f"Консультация для родителей: «Как поддержать этнокультурное развитие дома по «Кустук»»."
        }

        theme_k = st.text_input("Тема/Направление (если применимо)")
        if st.button("📝 Сгенерировать материал"):
            p_func = prompt_map[kustuk_section]
            with st.spinner("Генерация..."):
                res = ask_ai(p_func(theme_k or kustuk_section), ai_model, age_group, True)
            st.session_state.last_result = res
            st.markdown(res)
            show_exports(res, "Кустук")
            st.divider()
            if st.button("🖼️ Иллюстрация к материалу"):
                with st.spinner("Создаю изображение..."):
                    st.session_state.last_img = generate_image_shevdevrum(f"якутская культура, дети, {kustuk_section}, {age_group}, этностиль, акварель, образовательный постер")
            if st.session_state.last_img:
                st.image(st.session_state.last_img, caption=f"Иллюстрация: {kustuk_section}", use_container_width=True)

    elif page == "📖 НПА":
        st.info("📚 База нормативов в разработке. ИИ работает с актуальными версиями ФОП ДО, ФГОС ДО, СанПиН 2.4.3648-20, 273-ФЗ и методическими рекомендациями.")

    elif page == " Пользователи" and user["role"]=="admin":
        st.warning("🛡️ Управление аккаунтами. В MVP редактируйте `data/users.json` вручную.")
        st.code('{"admin@dou.ru": {"password":"bcrypt_hash", "role":"admin", "name":"Заведующая"}}', language="json")
